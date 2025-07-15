import io
from typing import Optional
import streamlit as st

# Import libraries from your requirements.txt
from docx import Document
from pdfminer.high_level import extract_text
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes

class ResumeParser:
    """Reads resume content from DOCX, PDF, or image files."""

    def read_file(self, uploaded_file) -> Optional[str]:
        """Dispatcher to read file content based on file type."""
        try:
            # Clone the file object so it can be read multiple times
            file_clone = io.BytesIO(uploaded_file.getvalue())
            
            filename = uploaded_file.name.lower()
            if filename.endswith('.docx'):
                # Call the new hybrid DOCX reader
                return self._read_docx_hybrid(file_clone)
            elif filename.endswith('.pdf'):
                return self._read_pdf_hybrid(file_clone)
            elif filename.endswith(('.png', '.jpg', '.jpeg')):
                return self._read_image(file_clone)
            else:
                st.error("Unsupported file format. Please upload a DOCX, PDF, or image.")
                return None
        except Exception as e:
            st.error(f"Error reading file '{uploaded_file.name}': {e}")
            return None

    def _read_docx_hybrid(self, file_bytes_io) -> str:
        """
        Extracts text from a DOCX file using a hybrid approach.
        1. Tries fast text extraction first.
        2. If that fails, falls back to OCR on any embedded images.
        """
        # --- Attempt 1: Fast Text Extraction ---
        try:
            doc = Document(file_bytes_io)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            text = "\n".join(full_text)
            if text and len(text.strip()) > 100:
                st.info("✅ Extracted text directly from DOCX.")
                return text
        except Exception as e:
            st.warning(f"Direct DOCX text extraction failed: {e}. Trying OCR fallback.")

        # --- Attempt 2: OCR Fallback ---
        st.info("ℹ️ Direct DOCX extraction yielded little text. Checking for images and using OCR...")
        try:
            # Rewind the BytesIO object to be read again
            file_bytes_io.seek(0)
            doc = Document(file_bytes_io)
            ocr_text = []
            
            # Iterate through the document's relationships to find all images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    image = Image.open(io.BytesIO(image_bytes))
                    ocr_text.append(pytesseract.image_to_string(image))
            
            if ocr_text:
                st.success("✅ Extracted text from images inside the DOCX file using OCR.")
                return "\n".join(ocr_text)
            else:
                st.warning("No text or images could be extracted from this DOCX file.")
                return "" # Return empty string if nothing is found
        except Exception as e:
            st.error(f"OCR processing on DOCX failed: {e}")
            return None

    def _read_pdf_hybrid(self, file_bytes_io) -> str:
        """Extracts text from a PDF using a hybrid approach."""
        # --- Attempt 1: Fast Text Extraction ---
        try:
            text = extract_text(file_bytes_io)
            if text and len(text.strip()) > 100:
                st.info("✅ Extracted text directly from PDF.")
                return text
        except Exception:
            pass # Silently fail and move to OCR

        # --- Attempt 2: OCR Fallback ---
        st.info("ℹ️ Direct PDF extraction failed or yielded little text. Falling back to OCR...")
        try:
            file_bytes_io.seek(0) # Rewind buffer
            images = convert_from_bytes(file_bytes_io.read())
            full_ocr_text = ""
            for i, image in enumerate(images):
                st.write(f"Processing page {i+1} with OCR...")
                full_ocr_text += pytesseract.image_to_string(image) + "\n"
            
            st.success("✅ Extracted text successfully using OCR.")
            return full_ocr_text
        except Exception as e:
            st.error(f"PDF OCR processing failed: {e}")
            return None

    def _read_image(self, file_bytes_io) -> str:
        """Extracts text from a direct image upload using Tesseract OCR."""
        image = Image.open(file_bytes_io)
        return pytesseract.image_to_string(image)
