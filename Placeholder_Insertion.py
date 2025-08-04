import io
from copy import deepcopy
from typing import Optional

import streamlit as st
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run


# Helper functions (unchanged)
def replace_text_in_paragraph(paragraph: Paragraph, key: str, value: str):
    """Replaces a specific key with a value within a paragraph's runs."""
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, str(value))

def copy_run_formatting(source_run: Run, target_run: Run):
    """Copies formatting from a source run to a target run."""
    if source_run is None or target_run is None:
        return
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.bold = source_run.font.bold
    target_run.font.italic = source_run.font.italic
    target_run.font.underline = source_run.font.underline
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb

def copy_paragraph_formatting(source_paragraph: Paragraph, target_paragraph: Paragraph):
    """Copies formatting from a source paragraph to a target paragraph."""
    target_paragraph.style = source_paragraph.style
    p_format = source_paragraph.paragraph_format
    target_p_format = target_paragraph.paragraph_format
    target_p_format.alignment = p_format.alignment
    target_p_format.space_after = p_format.space_after
    target_p_format.space_before = p_format.space_before
    target_p_format.left_indent = p_format.left_indent
    target_p_format.right_indent = p_format.right_indent

def replace_with_multiline_text(paragraph: Paragraph, key: str, formatted_text: str):
    """Replaces a placeholder with multiline text, preserving run formatting."""
    lines = formatted_text.split('\n')
    template_run = None
    for run in paragraph.runs:
        if key in run.text:
            template_run = run
            break
    if not template_run:
        return

    # Replace the key with the first line
    template_run.text = template_run.text.replace(key, lines[0])

    # Add subsequent lines with line breaks
    for line in lines[1:]:
        template_run.add_break()
        template_run.text += line

def replace_with_bullet_points(paragraph: Paragraph, key: str, bullet_points: list):
    """
    Replaces a placeholder with bullet points, creating new paragraphs for each.
    Also adds an "Achievements:" header if the key is {ACHIEVEMENTS}.
    """
    if not bullet_points:
        # If no bullet points, replace the placeholder with an empty string
        replace_text_in_paragraph(paragraph, key, "")
        return

    template_run = None
    for run in paragraph.runs:
        if key in run.text:
            template_run = run
            break
    if not template_run:
        return

    parent = paragraph._parent # Get the parent element (e.g., table cell)

    # Special handling for Achievements header
    if key == "{ACHIEVEMENTS}":
        header_p = parent.add_paragraph()
        copy_paragraph_formatting(paragraph, header_p) # Copy original paragraph formatting
        header_run = header_p.add_run("Achievements:")
        copy_run_formatting(template_run, header_run) # Copy original run formatting
        header_run.bold = True # Make the header bold

    # Remove the original paragraph containing the placeholder
    paragraph._element.getparent().remove(paragraph._element)

    # Add new paragraphs for each bullet point
    for point in bullet_points:
        new_p = parent.add_paragraph()
        copy_paragraph_formatting(paragraph, new_p) # Copy original paragraph formatting
        new_run = new_p.add_run("• " + point) # Add bullet character
        copy_run_formatting(template_run, new_run) # Copy original run formatting

    # Add an empty paragraph after achievements for spacing, if achievements were present
    if key == "{ACHIEVEMENTS}" and bullet_points:
        parent.add_paragraph("")


# Main function, adapted for Streamlit
def generate_resume(data: dict, template_path: str) -> Optional[io.BytesIO]:
    """
    Generates a resume by populating a .docx template with JSON data.
    This version returns an in-memory buffer for Streamlit.
    Handles None values in simple text fields by replacing them with "".
    """
    try:
        doc = Document(template_path)
    except Exception as e:
        st.error(f"Error opening template file: {e}")
        return None

    # 1. Simple replacements
    # Fetch values without a default, then handle None explicitly
    simple_replacements_data = {
        "{NAME}": data.get("name"),
        "{CONTACT}": data.get("contact_number"),
        "{EMAIL}": data.get("email"),
        "{NATIONALITY}": data.get("nationality"),
        "{SUMMARY}": data.get("summary"),
    }

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in simple_replacements_data.items():
                        if key in paragraph.text:
                            # Convert None to "" explicitly for display
                            display_value = "" if value is None else str(value)
                            replace_text_in_paragraph(paragraph, key, display_value)

    # 2. Education (Using row duplication as in your first script)
    education_placeholders = {
        "{DEGREE}", "{INSTITUTION}", "{EDUYEAR}", "{CGPA}"
    }

    for table in doc.tables:
        template_rows_info = []
        for i, row in enumerate(table.rows):
            row_text = "".join(cell.text for cell in row.cells)
            if any(ph in row_text for ph in education_placeholders):
                template_rows_info.append({'row': row, 'index': i})

        if not template_rows_info:
            continue

        template_rows = [info['row'] for info in template_rows_info]

        # Duplicate template rows for each education entry
        for experience in data.get("education", []):
            for template_row in template_rows:
                new_row_elem = deepcopy(template_row._element)
                table._tbl.append(new_row_elem) # Append duplicated row element
                new_row = table.rows[-1] # Get the newly added row object

                # Populate the new row's cells
                for cell in new_row.cells:
                    for p in list(cell.paragraphs): # Iterate over a copy to allow modification
                        replace_text_in_paragraph(p, "{DEGREE}", experience.get("degree", ""))
                        replace_text_in_paragraph(p, "{INSTITUTION}", experience.get("institution", ""))
                        replace_text_in_paragraph(p, "{EDUYEAR}", experience.get("year", ""))
                        replace_text_in_paragraph(p, "{CGPA}", experience.get("cgpa", ""))
        
        # Remove original template rows after duplication
        for template_row in reversed(template_rows): # Iterate in reverse to avoid index issues
            tbl = table._tbl
            tbl.remove(template_row._tr)
        
        # Assume only one education template block per table, break after processing
        break

    # 3. Skills and Languages (unchanged)
    list_replacements = {
        "{SKILLS}": data.get("skills", []),
        "{LANGUAGES}": data.get("languages", [])
    }

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    for key, value_list in list_replacements.items():
                        if key in p.text:
                            replace_with_bullet_points(p, key, value_list)

    # 4. Work Experience (unchanged)
    work_exp_placeholders = {
        "{COMPANYNAME}", "{DURATION}", "{JOBTITLE}",
        "{JOBDESCRIPTION}", "{ACHIEVEMENTS}"
    }

    for table in doc.tables:
        template_rows_info = []
        for i, row in enumerate(table.rows):
            row_text = "".join(cell.text for cell in row.cells)
            if any(ph in row_text for ph in work_exp_placeholders):
                template_rows_info.append({'row': row, 'index': i})

        if not template_rows_info:
            continue

        template_rows = [info['row'] for info in template_rows_info]

        for experience in data.get("work_experience", []):
            for template_row in template_rows:
                new_row_elem = deepcopy(template_row._element)
                table._tbl.append(new_row_elem)
                new_row = table.rows[-1]

                for cell in new_row.cells:
                    for p in list(cell.paragraphs):
                        replace_text_in_paragraph(p, "{COMPANYNAME}", experience.get("company_name", ""))
                        replace_text_in_paragraph(p, "{DURATION}", experience.get("duration", ""))
                        replace_text_in_paragraph(p, "{JOBTITLE}", experience.get("job_title", ""))

                        if "{JOBDESCRIPTION}" in p.text:
                            jd_list = experience.get("job_description", [])
                            replace_with_bullet_points(p, "{JOBDESCRIPTION}", jd_list)
                        
                        if "{ACHIEVEMENTS}" in p.text:
                            achievements_list = experience.get("achievements", [])
                            replace_with_bullet_points(p, "{ACHIEVEMENTS}", achievements_list)

        for template_row in reversed(template_rows):
            tbl = table._tbl
            tbl.remove(template_row._tr)

        # Assume only one work experience template block per table, break after processing
        break

    # Save to an in-memory buffer for Streamlit
    try:
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer # Return the buffer for the download button
    except Exception as e:
        st.error(f"Error saving output file to memory: {e}")
        return None
